from qfluentwidgets import RoundMenu, Action, FluentIcon, MenuAnimationType, MessageBox, Dialog, InfoBadge, InfoBadgePosition
from PyQt5.QtGui import QCursor
from PyQt5.QtCore import QPoint, QThread, pyqtSignal, QTimer
from PyQt5.QtWidgets import QLineEdit, QFileDialog
from ..models import StudentModel, Student, MouvementModel, \
    Mouvement, ComportementModel, Comportement, TypeComportementModel
from ..common import Function, col
from ..components import TableView
from ..view.students.list_student_tab import ListStudent
from ..view.students.new_student_dialog import NewStudentDialog
from ..view.students.show_student_dialog import ShowStudentDialog
from ..view.students.new_movement_dialog import NewMouvementDialog
from ..view.home.dialog.new_comp_dialog import NewComportementDialog
from ..view.students.import_dialog import ImportDialog
from .student_data_thread import DataThread
import os
from docx import Document

class StudentPresenter:
    
    HEADER_LABEL = ["Matricule", "Grade", "Nom", "Prénom(s)", "Genre", "Date de naissance", "Lieu de naissance"]
    
    def __init__(self, view:ListStudent, model: StudentModel, promotion):
        self.view = view
        self.model = model
        self.modelMove = MouvementModel()
        self.compModel = ComportementModel()
        self.typeCompModel = TypeComportementModel()
        self.promotion = promotion
        self.timer = QTimer()
        self.func = Function()
        self.view.tableView.setHorizontalHeaderLabels(self.HEADER_LABEL)
        self.__init_combox_data()
        self.actions()
        self.data = []
        self.fetchData()
        self.company = 0
        self.section = 0
        
    '''def updateLabel(self):
        comportement = self.compModel.fetch_items_by_id(self.promotion.id)
        for comp in comportement:
            self.HEADER_LABEL.append(comp.abrv)
            print(len(comportement))
            print(len(self.HEADER_LABEL))
        print(self.HEADER_LABEL)'''
        
    def __init_combox_data(self):
        companyStudents = self.model.fetch_items_by_id(self.promotion.id, group_by="company")
        sectionStudents = self.model.fetch_items_by_id(self.promotion.id, group_by="section")
        self.setComboxData(self.view.comboBoxCompany, companyStudents, "Compagnie", "company")
        self.setComboxData(self.view.comboBoxSection, sectionStudents, "Section", "section")
        
    def setComboxData(self, combox, data, label, key):
        nComboxData = []
        for item in data:
            val = item.get(key)
            companyLabel = f"{val}ère {label}" if val == 1 else f"{val}ème {label}"
            nComboxData.append(companyLabel)
        nComboxData.insert(0, label)
        combox.addItems(nComboxData)
        combox.setCurrentIndex(0)
        
    def actions(self):
        self.view.addAction.triggered.connect(lambda: self.addStudent())
        self.view.addComp.triggered.connect(lambda: self.addComp())
        self.view.importAction.triggered.connect(lambda: self.importData())
        self.view.exportAction.triggered.connect(self.export_data)
        self.view.deleteAction.triggered.connect(self.delete_data)
        self.view.searchLineEdit.textChanged.connect(self.on_text_changed)
        self.timer.setInterval(500)
        self.timer.timeout.connect(self.searchStudent)
        self.view.comboBoxCompany.currentTextChanged.connect(self.textChangedCombox)
        self.view.comboBoxSection.currentTextChanged.connect(self.textChangedCombox)
        self.view.toggleSelection.clicked.connect(self.reset)
        self.view.tableView.contextMenuEvent = lambda event: self.mouseRightClick(event)
        
    def reset(self):
        self.view.comboBoxCompany.setCurrentIndex(0)
        self.view.comboBoxSection.setCurrentIndex(0)
        self.view.searchLineEdit.setText("")
        
    def export_data(self):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getSaveFileName(self.view,"Exporter",f"{os.path.expanduser('~')}","CSV File (*.csv)", options=options)
        self.data.insert(0, self.HEADER_LABEL)
        if fileName:
            content_csv = ""
            for item in self.data:
                line = ""
                for val in item:
                    line += f'{val};'
                line += "\n"
                content_csv += line
            with open(fileName, 'w') as f:
                f.write(content_csv)
            os.startfile(fileName)
    
    def delete_data(self):
        w = Dialog("Supprimer", "Vous êtes sur de supprimer toutes les élèves, Toutes ces données seront suprrimées aussi!")
        w.setTitleBarVisible(False)
        if w.exec():
            self.view.parent.mainWindow.promModel.delete_all_student(self.promotion.id)
            self.fetchData()
        
    def on_text_changed(self):
        # Restart the timer whenever text is changed
        self.view.progressBar.setVisible(True)
        self.timer.start()
        
    def dataStudentFromDialog(self, dialog):
        lastname = dialog.lastnameEdit.lineEdit(0).text()
        firstname = dialog.firstnameEdit.lineEdit(0).text()
        matricule = dialog.matriculeEdit.lineEdit(0).text()
        level = dialog.gradeEdit.value()
        gender = dialog.genderEdit.value()
        return Student(
            promotion_id=self.promotion.id,
            lastname=lastname,
            firstname=firstname,
            gender=gender,
            level=level,
            matricule=matricule,
            company=matricule[0],
            section=matricule[1],
            number=matricule[2:4]
        )
            
    def addComp(self):
        dialog = NewComportementDialog(self.view)
        typeComp = self.typeCompModel.fetch_items_by_id(0)
        dataCombox = []
        for val in typeComp:
            dataCombox.append(val.name)
        dialog.typeCombox.addItems(dataCombox)
        
        if dialog.exec():
            name = dialog.nameLineEdit.text()
            abr = dialog.abrLineEdit.text()
            self.compModel.create(
                Comportement(
                    promotion_id=self.promotion.id,
                    name=name,
                    abrv=abr))
    
    def addStudent(self):
        dialog = NewStudentDialog(self.view.parent)
        if dialog.exec():
            student = self.dataStudentFromDialog(dialog)
            if self.model.count_by(matricule=student.matricule) == 0:
                self.model.create(student)
                self.fetchData()
            else:
                self.func.errorSuccess("Matricule invalide", "Matricule exist déjà", self.view.parent)
                
        
    def textChangedCombox(self, text:str):
        label_company = "compagnie"
        label_section = "section"
        if text.lower() != label_company and text.lower() != label_section:
            findLabel = text.split(" ")[1]
            value = text[0:text.find("è")]
            if findLabel.lower() == label_company:
                self.company = value
            else:
                self.section = value
        else:
            if text.lower() == label_company:
                self.company = 0
            else:
                self.section = 0
        query = {}
        if self.company != 0 and self.section != 0:
            query = {"section":self.section, "company":self.company}
        else:
            if self.company == 0:
                query = {"section":self.section}
            if self.section == 0:
                query = {"company": self.company}
                  
        if self.section == 0 and self.company == 0:
            self.fetchData()
        else:
            data = self.model.fetch_items_by_col(self.promotion.id, **query)
            self.worker_thread = DataThread(data, self.modelMove)
            self.worker_thread.update_progress.connect(self.update_progress_bar)
            self.worker_thread.finished.connect(self.worker_thread_finished)
            self.worker_thread.start()
    def mouseRightClick(self, event):
        selectedItems = self.view.tableView.selectedItems()
        if (len(selectedItems) != 0):
            matricule_item = self.view.tableView.selectedItems()[0].text()
            action = MenuAction(self)
            menu = RoundMenu(parent=self.view)
            menu.addAction(Action(FluentIcon.FOLDER, 'Voir', triggered = lambda:action.show(matricule_item)))
            menu.addAction(Action(FluentIcon.EDIT, 'Modifier', triggered = lambda: action.update(matricule_item)))
            menu.addSeparator()
            menu.addAction(Action(FluentIcon.SCROLL, 'Mouvement', triggered = lambda: action.mouvement(matricule_item)))
            menu.addSeparator()
            menu.addAction(Action(FluentIcon.DELETE, 'Supprimer', triggered = lambda: action.delete(matricule_item)))
            menu.menuActions()[-2].setCheckable(True)
            menu.menuActions()[-2].setChecked(True)

            self.posCur = QCursor().pos()
            cur_x = self.posCur.x()
            cur_y = self.posCur.y()
            menu.exec(QPoint(cur_x, cur_y), aniType=MenuAnimationType.FADE_IN_DROP_DOWN)
        
    def fetchData(self):
        data = self.model.fetch_items_by_id(self.promotion.id, order="matricule ASC")
        #self.view.tableView.setData(self.formatDataForTable(data))
        
        self.worker_thread = DataThread(data, self.modelMove)
        self.worker_thread.update_progress.connect(self.update_progress_bar)
        self.worker_thread.finished.connect(self.worker_thread_finished)
        self.worker_thread.start()
        
    def update_progress_bar(self, value):
        if value == 1:
            self.view.progressBar.setVisible(True)
        #self.view.progressBar.setValue(value)
        
    def worker_thread_finished(self):
        #self.view.progressBar.setValue(0)
        self.view.tableView.setData(self.worker_thread.listStudent)
        self.data = self.worker_thread.listStudent
        self.view.progressBar.setVisible(False)
        #print("ok")
        
    def formatDataForTable(self, data):
        listStudent = []
        permission = NewMouvementDialog.typesMove[0]
        rm = NewMouvementDialog.typesMove[1]
        exant = NewMouvementDialog.subType[0][1]
        anm = NewMouvementDialog.typesMove[3]
        codis = NewMouvementDialog.subType[1][0]
        horsTour = NewMouvementDialog.subType[1][1]
        bemolenge = NewMouvementDialog.subType[1][2]
        pertEfPol = NewMouvementDialog.subType[1][3]
        sanc_disc = NewMouvementDialog.typesMove[2]
        other = NewMouvementDialog.subType[1][4]
        lettreFel = NewMouvementDialog.subType[2][0]
        remarkPos = NewMouvementDialog.typesMove[4]
        for student in data:
            rmVal = self.modelMove.sum_by_with_id(student.id,"day",type=rm)
            exantVal = self.modelMove.sum_by_with_id(student.id,"day", subType=exant)
            permissionVal = self.modelMove.sum_by_with_id(student.id,"day",type=permission)
            anmVal = self.modelMove.sum_by_with_id(student.id,"day",type=anm)
            codisVal = self.modelMove.count_by_with_id(student.id, subType=codis)
            horsTourVal = self.modelMove.sum_by_with_id(student.id,"day",subType=horsTour)
            bemolengeVal = self.modelMove.count_by_with_id(student.id ,subType=bemolenge)
            pertEfPolVal = self.modelMove.count_by_with_id(student.id, subType=pertEfPol)
            other_sanct = self.modelMove.count_by_with_id(student.id, type=sanc_disc, subType=other)
            lettreFelVal = self.modelMove.count_by_with_id(student.id, subType=lettreFel)
            remarkPosVal = self.modelMove.count_by_with_id(student.id, type=remarkPos, subType=other)
            listStudent.append([
                student.matricule, student.level, student.lastname,
                student.firstname, student.gender,
                rmVal,exantVal,permissionVal,codisVal,bemolengeVal,
                horsTourVal,pertEfPolVal , other_sanct,anmVal,
                lettreFelVal,remarkPosVal])
        return listStudent
        
    def searchStudent(self):
        text = self.view.searchLineEdit.text()
        self.timer.stop()
        if len(text) > 2:
            data = self.model.search_with_id(self.promotion.id, matricule=text, firstname=text, lastname=text)
        #self.view.tableView.setData(self.formatDataForTable(data))
            self.worker_thread = DataThread(data, self.modelMove)
            self.worker_thread.update_progress.connect(self.update_progress_bar)
            self.worker_thread.finished.connect(self.worker_thread_finished)
            self.worker_thread.start()
        else:
            self.fetchData()
    
    def importData(self):
        filename = self.func.importFile(self.view, "Importer base de données", "CSV File (*.csv)")
        if filename:
            with open(filename, "r") as data:
                listStudent = []
                firstdata = []
                all_item = []
                keys = []
                for key in col.keys():
                    keys.append(key)
                for i, line in enumerate(data):
                    if i == 0:
                        firstdata = line.split(";")
                    else:
                        items = line.strip().split(";")
                        all_item.append(items)
                        
                dialog = ImportDialog(firstdata, Student(), self.view)
                if dialog.exec():
                    valCombox = []
                    index_combox = []
                    for i, comb in enumerate(dialog.combox):
                        index = comb.currentIndex() - 1
                        if index != -1:
                            valCombox.append(keys[index])
                        else:
                            valCombox.append('-')
                        index_combox.append(i)
                    for item in all_item:
                        entity = {"promotion_id" : self.promotion.id}
                        for i, val_combox in enumerate(valCombox):
                            if val_combox != "-":
                                if i < len(item):
                                    val = item[index_combox[i]]
                                    entity[val_combox] = val
                                    if val_combox == "matricule":
                                        if len(val) == 4:
                                            entity['company'] = val[0]
                                            entity['section'] = val[1]
                                            entity['number'] = val[2:]
                        if len(entity.keys()) > 1:
                            listStudent.append(Student(**entity))
                    if len(listStudent) > 0:
                        self.model.create_multiple(listStudent)
                        self.fetchData()
                    else:
                        self.func.errorMessage("Erreur", "Vous n'avez pas choisi même pas une seule colonne pour les données", self.view)
                
class MenuAction:
    
    def __init__(self,presenter:Student) -> None:
        self.model = presenter.model
        self.modelMove = presenter.modelMove
        self.view = presenter.view
        self.presenter = presenter
        
    def dataStudent(self, matricule):
        return self.model.fetch_item_by_cols(promotion_id=self.presenter.promotion.id, matricule=matricule)
       
    def show(self, matricule):
        student = self.dataStudent(matricule)
        dataStudent = f'{student.level} {student.matricule}\n{student.lastname} {student.firstname}'
        mouvements = self.presenter.modelMove.fetch_items_by_id(student.id)
        dataMouvements = []
        for mouvement in mouvements:
            dataMouvements.append([
                mouvement.date,
                f'{mouvement.type} - {mouvement.subType}' if len(mouvement.subType) != 0 and mouvement.subType != "-" else mouvement.type,
                mouvement.day
            ])
        dialog = ShowStudentDialog(self.view.parent.mainWindow)
        dialog.label.setText(dataStudent)
        if len(dataMouvements) > 0:
            total = int(self.presenter.modelMove.sum_by_with_id(student.id, "day"))
            if total > 0:
                dataMouvements.append(["Total", "", total])
            dialog.ImageMessage.setVisible(False)
            dialog.message.setVisible(False)
            dialog.table.setVisible(True)
            dialog.table.setData(dataMouvements)
        else:
            dialog.ImageMessage.setVisible(True)
            dialog.message.setVisible(True)
            dialog.table.setVisible(False)
        dialog.table.contextMenuEvent = lambda event, student=student, table=dialog.table: self.mouseRightClickTable(event, student, table)
        dialog.exportButton.clicked.connect(lambda: self.export_mouvement(student, mouvements))    
        dialog.exec()
        
    def mouseRightClickTable(self, event, student, table):
        if len(table.selectedItems()) != 0:
            date = table.selectedItems()[0].text()
            type_moves = table.selectedItems()[1].text().split(" - ")
            type_move = type_moves[0]
            sub_type_move = type_moves[1] if len(type_moves) > 1 else ""
            day = table.selectedItems()[2].text()
            menu = RoundMenu(parent=self.view)
            menu.addAction(Action(FluentIcon.DELETE, 'Supprimer', triggered=lambda:self.delete_move(student, date, type_move, sub_type_move, day, table)))
            self.posCur = QCursor().pos()
            cur_x = self.posCur.x()
            cur_y = self.posCur.y()

            menu.exec(QPoint(cur_x, cur_y), aniType=MenuAnimationType.FADE_IN_DROP_DOWN)
    
    def delete_move(self, student, date, type_m, sub_type_m, day, table):
        mod = {"date":date, "type":type_m, "day":day}
        if sub_type_m != "":
            mod["subType"] = sub_type_m
        w = Dialog("Supprimer", "Voulez-vous le supprimer vraiment", self.view.parent)
        w.setTitleBarVisible(False)
        if w.exec():
            self.modelMove.delete_by(idStudent=student.id, **mod)
            dataMouvements = []
            for mouvement in self.modelMove.fetch_items_by_id(student.id):
                dataMouvements.append([
                    mouvement.date,
                    f'{mouvement.type} - {mouvement.subType}' if len(mouvement.subType) != 0 and mouvement.subType != "-" else mouvement.type,
                    mouvement.day
                ])
            
            table.setData(dataMouvements)
            self.fetchData()
        
    def export_mouvement(self, student, mouvement):
        document = Document()
        document.add_heading(f'Informations', level=1)
        pData = f'Nom: {student.lastname}\n'
        pData += f'Prénoms: {student.firstname}\n'
        pData += f'Genre: {student.gender}\n'
        pData += f'Niveau: {student.level}\n'
        pData += f'Matricule: {student.matricule}\n'
        document.add_paragraph(pData)
        document.add_heading(f'Mouvements', level=1)
        
        # get table data -------------
        items = [[]]
        items.clear()
        day = "0"
        for mouv in mouvement:
            items.append([mouv.date, f"{mouv.type} {mouv.subType}", mouv.day])
            if(len(mouv.day) != 0):
                day +=  "+"+mouv.day
        valDay = eval(day)
        if valDay > 0:
            items.append(["Total", "",str(valDay)])

        if len(items) == 0:
            document.add_paragraph("Aucun mouvement")
        else:
            # add table ------------------
            table = document.add_table(1, 3)
            table.style = "Table Grid"

            # populate header row --------
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Date'
            heading_cells[1].text = 'Mouvement'
            heading_cells[2].text = 'Nombre de jour'

            # add a data row for each item
            for item in items:
                cells = table.add_row().cells
                cells[0].text = str(item[0])
                cells[1].text = item[1]
                cells[2].text = item[2]
    
        # Save the document
        filename = f"{os.path.expanduser('~')}"
        fileName = self.dialogSaveFile("Exporter", "", "Document Word (*.docx)")
        if fileName:
            document.save(fileName)
            os.startfile(fileName)  
        
    def dialogSaveFile(self, title:str, dir:str, typeFile:str):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getSaveFileName(self.view,title,dir,typeFile, options=options)
        return fileName    
    
    def mouvement(self, matricule):
        student = self.dataStudent(matricule)
        dialog = NewMouvementDialog(self.view.parent)
        dialog.subTitle.setText(f'{student.level} {student.lastname} {student.firstname}')
        if dialog.exec():
            mouvement = Mouvement(
                0,
                student.id,
                self.presenter.promotion.id,
                dialog.typeEdit.combox.text(),
                dialog.subTypeEdit.combox.text(),
                dialog.dateEdit.date.text(),
                dialog.dayMove.text())
            self.presenter.modelMove.create(mouvement)
            self.presenter.func.toastSuccess("Succès", "Ajout de mouvement de avec réussite", self.view.parent)
            self.fetchData()
         
    def delete(self, matricule):
        dialog = MessageBox(f"Supprimer", "Vous êtes sûr de vouloir supprimer?", self.view.parent.mainWindow)
        dialog.yesButton.setText("Oui")
        dialog.cancelButton.setText("Non")
        if dialog.exec():
            self.model.delete_by(promotion_id=self.presenter.promotion.id, matricule=matricule)
            self.fetchData()
            self.presenter.func.toastSuccess("Supprimé", "Suppression avec réussite", self.view.parent)
            
    def fetchData(self):
        text = self.view.searchLineEdit.text()
        if len(text) != 0:
            self.presenter.searchStudent()
        else:
            self.presenter.fetchData()
            
    def update(self, matricule):
        oldStudent = self.model.fetch_item_by_cols(promotion_id=self.presenter.promotion.id, matricule=matricule)
        dialog = NewStudentDialog(self.view.parent)
        dialog.lastnameEdit.lineEdit(0).setText(oldStudent.lastname)
        dialog.firstnameEdit.lineEdit(0).setText(oldStudent.firstname)
        dialog.matriculeEdit.lineEdit(0).setText(str(oldStudent.matricule))
        if oldStudent.gender == "F":
            dialog.genderEdit.combox.setCurrentIndex(1)
        if oldStudent.level == "EAP":
            dialog.gradeEdit.combox.setCurrentIndex(1)
            
        dialog.yesButton.setEnabled(True)
        dialog.yesButton.setText("Mettre à jour")
        if dialog.exec():
            student = self.presenter.dataStudentFromDialog(dialog)
            self.model.update_item(oldStudent.id, 
                              lastname=student.lastname, 
                              firstname=student.firstname,
                              gender=student.gender,
                              level=student.level,
                              matricule=student.matricule,
                              company=student.company,
                              section=student.section,
                              number=student.number)
            self.fetchData()
            self.presenter.func.toastSuccess("Mise à jour", "Mise à jour d'élève avec réussite!", self.view.parent)
       