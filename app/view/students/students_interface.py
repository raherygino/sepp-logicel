# coding:utf-8
from ..utils.gallery_interface import GalleryInterface
from ...common.translator import Translator
from ...common.Translate import Translate
from ...common.config import Lang
from ...common.keys import *
from ...components import *
from qfluentwidgets import (SubtitleLabel, SearchLineEdit, ToolButton,MenuAnimationType, 
                            ToggleToolButton,ComboBox, RoundMenu, Action, MessageBox, InfoBar, InfoBarPosition)
from qfluentwidgets import FluentIcon as FIF
from PyQt5.QtWidgets import QFrame, QVBoxLayout, QTableWidgetItem, QAction, QFileDialog
from PyQt5.QtCore import Qt, QSize, QCoreApplication, QModelIndex, QPoint
from PyQt5.QtGui import QCursor

from ...common.config import *
from .students_new_dialog import DialogStudent
from .student_show import DialogStudentShow
from .student_move import DialogStudentMove
from ...common.database.service.student_service import StudentService
from PyQt5.QtSql import QSqlDatabase
from ...common.database.db_initializer import DBInitializer as DB
from ...common.database.dao.student_dao import Student
from ...common.database.entity import Mouvement
from ...common.database.service.mouvement_service import MouvementService
from ...common.database.utils.constants import *
from docx import Document
import os

class StudentInterface(GalleryInterface):
    """ Student interface """

    def __init__(self, parent=None):
        self.db = parent.db
        self.trans = Translate(Lang().current).text
        self.companySelected = 0
        self.sectionSelected = 0
        self.isSelection = False

        super().__init__(
            title='',
            subtitle='',
            parent=parent
        )
        
        self.db = QSqlDatabase.database(self.db.CONNECTION_NAME, True)
        self.studentService = StudentService(self.db)
        self.moveService = MouvementService(self.db)
        self.parent = parent
        self.hBoxLayout = QVBoxLayout(self)
        self.titleContainte(parent)
        self.container(parent)
        self.setObjectName('studentInterface')

    def titleContainte(self, parent):
        row = Frame(VERTICAL, ROW+str(1), parent=parent)
        label = SubtitleLabel("Liste des élèves de la promotion Sandratra", parent)
        row.setMargins(9,0,9,0)
        row.addWidget(label)
        self.hBoxLayout.addWidget(row)

    def container(self, parent):
        self.container = Frame(VERTICAL, STUDENT+CONTAINER, parent=parent)
        self.row_2 = Frame(HORIZONTAL, ROW+str(2), parent=parent)
        self.searchLineStudent = SearchLineEdit(self)
        self.searchLineStudent.setPlaceholderText(QCoreApplication.translate(FORM, u"Recherche", None))
        self.searchLineStudent.setMaximumSize(QSize(240, 50))
        self.searchLineStudent.textChanged.connect(self.searchStudent)
        
        col = Frame(HORIZONTAL, COL+str(1),parent=parent)
        
        self.comboBoxCompany = ComboBox(self)

        companies = []
        for company in self.studentService.companies():
            if company == "1":
                companies.append(f'{company}ère Compagnie')
            else:
                companies.append(f'{company}ème Compagnie')

        companies.insert(0, "Companie")
        self.comboBoxCompany.addItems(companies)
        self.comboBoxCompany.setCurrentIndex(0)
        self.comboBoxCompany.currentTextChanged.connect(self.textChangedCombox)
        self.comboBoxCompany.move(200, 200)
        self.comboBoxSection = ComboBox(self)

        sections = []
        sections.clear()
        for section in self.studentService.sections():
            if section == "1":
                sections.append(f'{section}ère Section')
            else:
                sections.append(f'{section}ème Section')
        sections.insert(0, "Section")

        self.comboBoxSection.addItems(sections)
        self.comboBoxSection.setCurrentIndex(0)
        self.comboBoxSection.currentTextChanged.connect(self.textChangedCombox)
        self.comboBoxSection.move(200, 200)

        # toggle tool button
        self.toggleSelection = ToggleToolButton(FIF.FILTER, self)
        self.toggleSelection.toggled.connect(self.setSelection)
        #self.toggleSelection.toggle()
        
        # tool button
        self.toolButton = ToolButton(FIF.SAVE, self)
        self.toolButton.clicked.connect(self.exportExcel)

        self.comboBoxCompany.setEnabled(False)
        self.comboBoxSection.setEnabled(False)

        col.layout.addWidget(self.comboBoxCompany)
        col.layout.addWidget(self.comboBoxSection)
        col.layout.addWidget(self.toggleSelection)
        col.layout.addWidget(self.toolButton)

        col.setMargins(0,0,0,0)
        
        self.row_2.setMargins(0,0,0,0)
        self.row_2.addWidget(self.searchLineStudent)
        self.row_2.layout.addWidget(col, 0, Qt.AlignRight)
        self.container.addWidget(self.row_2)

        students = self.listStudent(self.studentService.listAll())
        self.tbStudent = Table(parent, students.get("header"), students.get("data"))
        self.tbStudent.setRisizeMode(len(students.get("header")) - 1)
        self.table = self.tbStudent.widget()
        self.table.clicked.connect(self.selectItem)
        self.container.addWidget(self.tbStudent.widget())
        self.hBoxLayout.addWidget(self.container)
        self.dialog = None

    def setSelection(self):
        if self.isSelection:
            if self.comboBoxCompany.currentIndex() != 0 and self.comboBoxSection.currentIndex() != 0:  
                self.refreshTable()
            self.comboBoxCompany.setCurrentIndex(0)
            self.comboBoxCompany.setEnabled(False)
            self.comboBoxSection.setCurrentIndex(0)
            self.comboBoxSection.setEnabled(False)
            self.toggleSelection.setIcon(FIF.FILTER)
            self.isSelection = False
        else:
            self.comboBoxCompany.setEnabled(True)
            self.comboBoxSection.setEnabled(True)
            self.toggleSelection.setIcon(FIF.CLOSE)
            self.isSelection = True

    def textChangedCombox(self, text:str):
        value = ""
        pos = text.find("è")
        name = text
        isSelected = pos != -1

        if isSelected:
            name = text.split(" ")[1]
            value = text[0:pos]

        if name == "Compagnie":
            self.companySelected = value

        elif name == "Section":
            self.sectionSelected = value

        data = self.studentService.listByFields(
            company=self.companySelected, 
            section=self.sectionSelected)
        
        if self.companySelected == "0" and self.sectionSelected == "0":
            data = self.studentService.listAll()

        listStudent = self.listStudent(data)
        self.tbStudent.refresh(self.table, listStudent.get("header"), listStudent.get("data"))
        
        #print(self.companySelected)
        #print(self.sectionSelected)

    def listStudent(self, data):
        stude = self.studentService

        RM = TYPE_MOVE["RM_CONV"]
        EX = SUB_TYPE_MOVE["EX_PHYS"]
        PERM = TYPE_MOVE["PERMISSION"]
        CODIS = SUB_TYPE_MOVE["CODIS"]
        HORS_TOUR = SUB_TYPE_MOVE["HORS_TOUR"]
        BEM = SUB_TYPE_MOVE["BEM"]
        PEP = SUB_TYPE_MOVE["PEP"]
        ANM = TYPE_MOVE["ANM"]
        LETTRE_FEL = SUB_TYPE_MOVE["LETTRE_FEL"]
        OTHER_SAC = f'{SUB_TYPE_MOVE["OTHER"]} {TYPE_MOVE["SACT_DISC"]}'
        OTHER_REM_POS = f'{SUB_TYPE_MOVE["OTHER"]} {TYPE_MOVE["REM_POS"]}'

        header = [
            "ID","Matricule", "Niveau", "Nom",
            "Prénom", "Genre",RM+" (Jour)",
            EX+" (Jour)", PERM+" (Jour)", CODIS+" (fois)", HORS_TOUR+" (Jour)",
            BEM+" (fois)", PEP+" (fois)", 
            OTHER_SAC+" (fois)", ANM+" (Jour)", 
            LETTRE_FEL+" (fois)", OTHER_REM_POS+" (fois)"]
        listStudent = [[]]
        listStudent.clear()

        for student in data:
            
            idStudent = student.get("id_tbl_student")
            listStudent.append([
                idStudent,
                student.get("matricule"),
                student.get("level"),
                student.get("lastname"),
                student.get("firstname"),
                student.get("gender"),
                stude.sumOfDayTypeMove(idStudent, RM), 
                stude.sumOfDaySubTypeMove(idStudent, EX),
                stude.sumOfDayTypeMove(idStudent, PERM),
                stude.countSubTypeMove(idStudent, CODIS),
                stude.sumOfDaySubTypeMove(idStudent, HORS_TOUR),
                stude.countSubTypeMove(idStudent, BEM),
                stude.countSubTypeMove(idStudent, PEP),
                stude.countSubTypeMove(idStudent, OTHER_SAC),
                stude.sumOfDayTypeMove(idStudent, ANM),
                stude.countSubTypeMove(idStudent, LETTRE_FEL),
                stude.countSubTypeMove(idStudent, OTHER_REM_POS)
                ]) 
        return {
            "header": header,
            "data": listStudent
        }
    
    def refreshTable(self, **kwargs):
        data = self.studentService.listAll()
        if (len(kwargs.keys()) != 0):
            data = self.studentService.search(kwargs.get('query'))

        listStudent = self.listStudent(data)
        self.tbStudent.refresh(self.table, listStudent.get("header"), listStudent.get("data"))

    def showDialog(self):
        self.dialog = DialogStudent(self.parent)
        self.dialog.title.setText("Add new student")
        btnOk = self.dialog.yesButton
        btnOk.setText("Create")
        btnOk.clicked.connect(lambda: self.createStudent(self.dialog.studentData()))
        self.dialog.show()


    def showDialogItem(self, id):
        self.dialog = DialogStudentShow(self.studentService, self.moveService, id, self.parent)
        self.dialog.yesButton.clicked.connect(lambda: self.exportData(self.dialog.student, self.dialog.d))
        self.dialog.show()
    
    def showDialogView(self, item:QModelIndex):
        id = self.table.item(item.row(), 0).text()
        self.dialog = DialogStudent(self.parent, service=self.studentService,id=id)
        self.dialog.title.setText(f"Update {self.dialog.student.lastname}")
        btnOk = self.dialog.yesButton
        btnOk.setText("Update")
        btnOk.clicked.connect(lambda: self.updateStudent(id, self.dialog.studentData()))
        self.dialog.show()

    def createStudent(self, student: Student):
        if (self.studentService.create(student)):
            self.infoMessage(self.parent, "Created", "Student created succesfully!")
            self.dialog.accept()
            self.dialog = None
            self.refreshTable()
        else:
            print("error")
    def updateStudent(self, id:int, student: Student):
            if(self.studentService.update(id, student)):
                self.infoMessage(self.parent, "Update", f"{student.firstname} updated succesfully!")
                self.dialog.accept()
                self.dialog = None
                self.refreshTable()
            else:
                print("Error")
    
    def searchStudent(self, text:str):
        if '\'' not in text:
            self.refreshTable(query=text)
            self.sectionSelected = 0
            self.companySelected = 0
            self.comboBoxCompany.setEnabled(False)
            self.comboBoxSection.setEnabled(False)
            self.toggleSelection.setIcon(FIF.FILTER)
            self.toggleSelection.setChecked(False)


    def exportExcel(self):
        data = ""
        listData = self.tbStudent.latestData
        
        for i, h in enumerate(self.tbStudent.h):
            if i > 0:
                data += f"{h};"
        data += "\n"

        for val in listData:
            line = ""
            for i, value in enumerate(val):
                if i > 0:
                    line += f"{value};"
            data += f"{line}\n"

        dir_recent = f"{os.path.expanduser('~')}\Documents"
        fileName = self.dialogSaveFile("Exporter", dir_recent, "CSV File (*.CSV)")

        if fileName:
            if fileName.find('.csv') == -1:
                fileName += '.csv'
            self.saveFile(data, fileName)
    
    def saveFile(self, data, filename):
        with open(filename, "w") as file:
            file.write(data)
        os.startfile(filename)


    def dialogSaveFile(self, title:str, dir:str, typeFile:str):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self,title,dir,typeFile, options=options)
        return fileName
      

     
    def dialogOpenFile(self):
        '''
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        dir_recent = f"{os.path.expanduser('~')}\Documents"
        fileName, _ = QFileDialog.getOpenFileName(self,"Importer", dir_recent,"CSV Files (*.xlsx);;All Files (*)", options=options)
        '''
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        data = 'hello'
        dir_recent = f"{os.path.expanduser('~')}\Documents"
        fileName, _ = QFileDialog.getSaveFileName(self,"Exporter",dir_recent+"\""+data[1],"CSV File (*.CSV)", options=options)
        if fileName:
            print("ok")
        
    def exportData(self, student: Student, mouvement):
        document = Document()
        document.add_heading(f'Informations', level=1)
        pData = f'Nom: {student.lastname}\n'
        pData += f'Prénoms: {student.firstname}\n'
        pData += f'Genre: {student.gender}\n'
        pData += f'Niveau: {student.level}\n'
        document.add_paragraph(pData)
        document.add_heading(f'Mouvements', level=1)
        
        # get table data -------------
        items = [[]]
        items.clear()
        day = "0"
        for mouv in mouvement:
            items.append([mouv.date, f"{mouv.type} {mouv.subType}", mouv.day])
            if(len(mouv.day) != 0):
                day +=  "+"+mouv.day
        valDay = eval(day)
        if valDay > 0:
            items.append(["Total", "",str(valDay)])

        if len(items) == 0:
            document.add_paragraph("Aucun mouvement")
        else:
            # add table ------------------
            table = document.add_table(1, 3)
            table.style = "Table Grid"

            # populate header row --------
            heading_cells = table.rows[0].cells
            heading_cells[0].text = 'Date'
            heading_cells[1].text = 'Mouvement'
            heading_cells[2].text = 'Nombre de jour'

            # add a data row for each item
            for item in items:
                cells = table.add_row().cells
                cells[0].text = str(item[0])
                cells[1].text = item[1]
                cells[2].text = item[2]
    
        # Save the document
        filename = f"{os.path.expanduser('~')}\Documents\{student.level}-{student.matricule}.docx";
        fileName = self.dialogSaveFile("Exporter", filename, "Document Word (*.docx)")
        if fileName:
            document.save(filename)
            self.dialog.yesButton.setVisible(False)
            os.startfile(filename)  
        

    
    def selectItem(self, item: QModelIndex):
        menu = RoundMenu(parent=self)
        menu.addAction(Action(FIF.FOLDER, 'Voir', triggered=lambda:self.showItem(item)))
        menu.addAction(Action(FIF.EDIT, 'Modifier', triggered=lambda:self.showDialogView(item)))
        menu.addSeparator()
        menu.addAction(Action(FIF.SCROLL, 'Mouvement', triggered=lambda:self.showDialogMove(item)))
       # menu.addAction(Action(FIF.DELETE, 'Delete', triggered=lambda:self.confirmDeleteItem(item)))
        menu.menuActions()[-2].setCheckable(True)
        menu.menuActions()[-2].setChecked(True)

        self.posCur = QCursor().pos()
        cur_x = self.posCur.x()
        cur_y = self.posCur.y()

        menu.exec(QPoint(cur_x, cur_y), aniType=MenuAnimationType.DROP_DOWN)

    def confirmDeleteItem(self, item: QModelIndex):
        confirm = MessageBox("Confirmation", "You want to delete it?", self.parent)
        confirm.accepted.connect(lambda:self.deleteItem(item))
        confirm.show()

    def showDialogMove(self, item: QModelIndex):
        id = self.table.item(item.row(), 0).text()
        self.dialog = DialogStudentMove(self.studentService, self.moveService, id, self.parent)
        btn = self.dialog.yesButton
        btn.clicked.connect(lambda: self.newMouvement(self.dialog.dataMouvement()))
        self.dialog.show()
    
    def newMouvement(self, mouvement:Mouvement):
        service = MouvementService(self.db)
        service.create(mouvement)
        searchQuery = self.searchLineStudent.text()
        if searchQuery != "":
            self.refreshTable(query=searchQuery)
        else:
            if self.companySelected != 0 and self.sectionSelected != 0:
                data = self.studentService.listByFields(
                    company=self.companySelected, 
                    section=self.sectionSelected)
                listStudent = self.listStudent(data)
                self.tbStudent.refresh(self.table, listStudent.get("header"), listStudent.get("data"))
            else:
                self.refreshTable()
        self.dialog.accept()
    
    def showItem(self, item: QModelIndex):
        id = self.table.item(item.row(), 0).text()
        self.showDialogItem(id)
        

    def deleteItem(self, item: QModelIndex):
        id = self.table.item(item.row(), 0).text()
        self.studentService.deleteById(id)
        self.infoMessage(self.parent, "Success", "Student deleted successfully")
        self.refreshTable()

    def infoMessage(self, parent, title:str, message:str):
        InfoBar.success(
            title=title,
            content=message,
            orient=Qt.Horizontal,
            isClosable=True,
            position=InfoBarPosition.TOP,
            duration=2000,
            parent=parent
        )